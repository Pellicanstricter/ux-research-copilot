# agents.py - Clean Working Version
from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Any, Tuple
import asyncio
import logging
from datetime import datetime
import json
import re
import hashlib
from pathlib import Path
import os

# External dependencies
try:
    import faiss
    import numpy as np
    from langchain_openai import OpenAIEmbeddings
except ImportError:
    print("⚠️  Some optional dependencies missing - continuing without them")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
import redis
import docx
import PyPDF2

from models import InsightData, ThemeCluster, PersonaData, ProcessingStatus, KeyInsightCard, ExecutiveSummary, QuoteWithAttribution
from config import AgentConfig

# Configure logging
logging.basicConfig(level=logging.INFO)

# Base Agent Interface
class BaseAgent(ABC):
    """Abstract base class for all agents in the system"""
    
    def __init__(self, config: AgentConfig, session_id: str):
        self.config = config
        self.session_id = session_id
        self.logger = logging.getLogger(f"{self.__class__.__name__}")
        
        try:
            self.redis_client = redis.Redis(
                host=config.redis_host, 
                port=config.redis_port, 
                decode_responses=True
            )
            # Test connection
            self.redis_client.ping()
        except redis.ConnectionError:
            self.logger.warning("Redis connection failed - running without session persistence")
            self.redis_client = None
        
    @abstractmethod
    async def process(self, input_data: Any) -> Any:
        """Process input data and return results"""
        pass
    
    def update_session_status(self, status: ProcessingStatus, **kwargs):
        """Update session status in Redis"""
        if not self.redis_client:
            return
            
        try:
            session_key = f"session:{self.session_id}"
            session_data = self.redis_client.hgetall(session_key)
            
            session_data.update({
                'status': status.value,
                'updated_at': datetime.now().isoformat(),
                **kwargs
            })
            
            self.redis_client.hset(session_key, mapping=session_data)
        except Exception as e:
            self.logger.warning(f"Failed to update session status: {e}")

# Agent 1: Document Ingestor
class DocumentIngestor(BaseAgent):
    """Processes and chunks transcript files with metadata extraction"""
    
    def __init__(self, config: AgentConfig, session_id: str):
        super().__init__(config, session_id)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            length_function=len
        )
    
    async def process(self, file_paths: List[str]) -> List[Document]:
        """Process multiple files and return chunked documents"""
        self.update_session_status(ProcessingStatus.PROCESSING)
        
        try:
            all_documents = []
            
            for file_path in file_paths:
                documents = await self._process_single_file(file_path)
                all_documents.extend(documents)
            
            self.logger.info(f"Processed {len(file_paths)} files into {len(all_documents)} chunks")
            return all_documents
            
        except Exception as e:
            self.logger.error(f"Error processing files: {str(e)}")
            self.update_session_status(ProcessingStatus.FAILED, error_message=str(e))
            raise
    
    async def _process_single_file(self, file_path: str) -> List[Document]:
        """Process a single file based on its type"""
        path = Path(file_path)
        
        try:
            if path.suffix.lower() == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif path.suffix.lower() in ['.docx', '.doc']:
                text = self._extract_from_docx(file_path)
            elif path.suffix.lower() == '.txt':
                text = self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {path.suffix}")
            
            # Clean and preprocess text
            text = self._preprocess_text(text)
            
            # Create base document with metadata
            base_doc = Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "file_type": path.suffix,
                    "processed_at": datetime.now().isoformat(),
                    "session_id": self.session_id
                }
            )
            
            # Split into chunks
            chunks = self.text_splitter.split_documents([base_doc])
            
            # Add chunk-specific metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata.update({
                    "chunk_id": i,
                    "total_chunks": len(chunks),
                    "chunk_hash": hashlib.md5(chunk.page_content.encode()).hexdigest()
                })
            
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {str(e)}")
            return []
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            self.logger.error(f"Error extracting PDF {file_path}: {e}")
            raise
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            self.logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
        return text
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            self.logger.error(f"Error extracting TXT {file_path}: {e}")
            raise
    
    def _preprocess_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:\-\'"()]', '', text)
        
        # Normalize speaker patterns
        text = re.sub(r'^([A-Za-z0-9\s]+):\s*', r'\1: ', text, flags=re.MULTILINE)
        
        return text.strip()

# Agent 2: Insight Analyzer
class InsightAnalyzer(BaseAgent):
    """Extracts key quotes and identifies themes using NLP"""
    
    def __init__(self, config: AgentConfig, session_id: str):
        super().__init__(config, session_id)
        
        if not config.openai_api_key:
            raise ValueError("OpenAI API key is required for InsightAnalyzer")
            
        self.llm = ChatOpenAI(
            api_key=config.openai_api_key,
            model="gpt-3.5-turbo",
            temperature=config.temperature,
            max_tokens=config.max_tokens
        )
        self.parser = JsonOutputParser(pydantic_object=InsightData)
        self._setup_prompts()
    
    def _setup_prompts(self):
        """Setup prompts for insight extraction"""
        self.insight_prompt = PromptTemplate(
            input_variables=["text_chunk"],
            template="""
            As an expert UX researcher, analyze the following transcript chunk and extract key insights.
            
            Transcript:
            {text_chunk}
            
            Extract insights focusing on:
            1. User pain points and frustrations
            2. User goals and motivations  
            3. Behavioral patterns
            4. Feature requests or suggestions
            5. Emotional reactions
            
            For each insight, provide:
            - Exact quote from the transcript
            - Speaker (if identifiable)
            - Theme category (e.g., "Navigation Issues", "Feature Request", "User Goal")
            - Sentiment (Positive, Negative, or Neutral)
            - Confidence score (0-1)
            - Context (brief description)
            - Timestamp (if available)
            
            Return a JSON list of insights. If no significant insights are found, return an empty list.
            
            {format_instructions}
            """
        ).partial(format_instructions=self.parser.get_format_instructions())

    async def process(self, documents: List[Document]) -> List[InsightData]:
        """Process documents and extract insights"""
        try:
            all_insights = []
            
            for doc in documents:
                chunk_insights = await self._analyze_chunk(doc.page_content)
                all_insights.extend(chunk_insights)
            
            # Remove duplicates
            unique_insights = self._deduplicate_insights(all_insights)
            
            self.logger.info(f"Extracted {len(unique_insights)} unique insights from {len(documents)} chunks")
            
            return unique_insights
            
        except Exception as e:
            self.logger.error(f"Error analyzing insights: {str(e)}")
            raise

    async def _analyze_chunk(self, chunk: str) -> List[InsightData]:
        """Analyze a single chunk and extract insights"""
        try:
            # Create the prompt for this chunk
            formatted_prompt = self.insight_prompt.format(text_chunk=chunk)
            
            # Make the API call
            response = await self.llm.ainvoke(formatted_prompt)
            
            # Parse the response
            try:
                insights_data = self.parser.parse(response.content)
                if isinstance(insights_data, list):
                    return [InsightData(**insight) for insight in insights_data]
                else:
                    return [InsightData(**insights_data)]
            except:
                # Fallback to simple parsing if JSON parsing fails
                return self._parse_insights_fallback(response.content)
            
        except Exception as e:
            self.logger.error(f"Error analyzing chunk: {str(e)}")
            return []

    def _parse_insights_fallback(self, response_text: str) -> List[InsightData]:
        """Fallback parsing when JSON parsing fails"""
        insights = []
        lines = response_text.strip().split('\n')
        
        current_insight = {}
        for line in lines:
            line = line.strip()
            if line.startswith('Quote:'):
                current_insight['quote'] = line.replace('Quote:', '').strip()
            elif line.startswith('Theme:'):
                current_insight['theme'] = line.replace('Theme:', '').strip()
            elif line.startswith('Sentiment:'):
                current_insight['sentiment'] = line.replace('Sentiment:', '').strip()
            elif line.startswith('Confidence:'):
                try:
                    current_insight['confidence'] = float(line.replace('Confidence:', '').strip())
                except:
                    current_insight['confidence'] = 0.5
            elif line.startswith('Context:'):
                current_insight['context'] = line.replace('Context:', '').strip()
                # Complete insight, add to list
                if 'quote' in current_insight and 'theme' in current_insight:
                    insight_data = InsightData(
                        quote=current_insight.get('quote', ''),
                        theme=current_insight.get('theme', 'General'),
                        sentiment=current_insight.get('sentiment', 'Neutral'),
                        confidence=current_insight.get('confidence', 0.5),
                        context=current_insight.get('context', ''),
                        speaker=current_insight.get('speaker'),
                        timestamp=current_insight.get('timestamp')
                    )
                    insights.append(insight_data)
                current_insight = {}
        
        return insights# Agent 2: Insight Analyzer
class InsightAnalyzer(BaseAgent):
    """Extracts key quotes and identifies themes using NLP"""
    
    def __init__(self, config: AgentConfig, session_id: str):
        super().__init__(config, session_id)
        
        if not config.openai_api_key:
            raise ValueError("OpenAI API key is required for InsightAnalyzer")
            
        self.llm = ChatOpenAI(
            api_key=config.openai_api_key,
            model="gpt-3.5-turbo",
            temperature=config.temperature,
            max_tokens=config.max_tokens
        )
        self.parser = JsonOutputParser(pydantic_object=InsightData)
        self._setup_prompts()
    
    def _setup_prompts(self):
        """Setup prompts for insight extraction"""
        self.insight_prompt = PromptTemplate(
            input_variables=["text_chunk"],
            template="""
            As an expert UX researcher, analyze the following transcript chunk and extract key insights.
            
            Transcript:
            {text_chunk}
            
            Extract insights focusing on:
            1. User pain points and frustrations
            2. User goals and motivations  
            3. Behavioral patterns
            4. Feature requests or suggestions
            5. Emotional reactions
            
            For each insight, provide:
            - Exact quote from the transcript
            - Speaker (if identifiable)
            - Theme category (e.g., "Navigation Issues", "Feature Request", "User Goal")
            - Sentiment (Positive, Negative, or Neutral)
            - Confidence score (0-1)
            - Context (brief description)
            - Timestamp (if available)
            
            Return a JSON list of insights. If no significant insights are found, return an empty list.
            
            {format_instructions}
            """
        ).partial(format_instructions=self.parser.get_format_instructions())

    async def process(self, documents: List[Document]) -> List[InsightData]:
        """Process documents and extract insights"""
        try:
            all_insights = []
            
            for doc in documents:
                chunk_insights = await self._analyze_chunk(doc.page_content)
                all_insights.extend(chunk_insights)
            
            # Remove duplicates
            unique_insights = self._deduplicate_insights(all_insights)
            
            self.logger.info(f"Extracted {len(unique_insights)} unique insights from {len(documents)} chunks")
            
            return unique_insights
            
        except Exception as e:
            self.logger.error(f"Error analyzing insights: {str(e)}")
            raise

    async def _analyze_chunk(self, chunk: str) -> List[InsightData]:
        """Analyze a single chunk and extract insights"""
        try:
            # Create the prompt for this chunk
            formatted_prompt = self.insight_prompt.format(text_chunk=chunk)
            
            # Make the API call
            response = await self.llm.ainvoke(formatted_prompt)
            
            # Parse the response
            try:
                insights_data = self.parser.parse(response.content)
                if isinstance(insights_data, list):
                    return [InsightData(**insight) for insight in insights_data]
                else:
                    return [InsightData(**insights_data)]
            except Exception as e:
                # Log the parsing error and the actual response
                self.logger.error(f"JSON parsing failed: {str(e)}")
                self.logger.error(f"Raw response content: {response.content[:500]}...")
                
                # Fallback to simple parsing if JSON parsing fails
                fallback_results = self._parse_insights_fallback(response.content)
                self.logger.info(f"Fallback parser returned {len(fallback_results)} insights")
                return fallback_results
            
        except Exception as e:
            self.logger.error(f"Error analyzing chunk: {str(e)}")
            return []

    def _parse_insights_fallback(self, response_text: str) -> List[InsightData]:
        """Fallback parsing when JSON parsing fails"""
        insights = []
        lines = response_text.strip().split('\n')
        
        current_insight = {}
        for line in lines:
            line = line.strip()
            if line.startswith('Quote:'):
                current_insight['quote'] = line.replace('Quote:', '').strip()
            elif line.startswith('Theme:'):
                current_insight['theme'] = line.replace('Theme:', '').strip()
            elif line.startswith('Sentiment:'):
                current_insight['sentiment'] = line.replace('Sentiment:', '').strip()
            elif line.startswith('Confidence:'):
                try:
                    current_insight['confidence'] = float(line.replace('Confidence:', '').strip())
                except:
                    current_insight['confidence'] = 0.5
            elif line.startswith('Context:'):
                current_insight['context'] = line.replace('Context:', '').strip()
                # Complete insight, add to list
                if 'quote' in current_insight and 'theme' in current_insight:
                    insight_data = InsightData(
                        quote=current_insight.get('quote', ''),
                        theme=current_insight.get('theme', 'General'),
                        sentiment=current_insight.get('sentiment', 'Neutral'),
                        confidence=current_insight.get('confidence', 0.5),
                        context=current_insight.get('context', ''),
                        speaker=current_insight.get('speaker'),
                        timestamp=current_insight.get('timestamp')
                    )
                    insights.append(insight_data)
                current_insight = {}
        
        return insights

    def _deduplicate_insights(self, insights: List[InsightData]) -> List[InsightData]:
        """Remove duplicate insights based on quote similarity"""
        unique_insights = []
        seen_quotes = set()
        
        for insight in insights:
            # Simple deduplication based on quote content
            quote_key = insight.quote.lower().strip()[:50]  # First 50 chars, normalized
            if quote_key not in seen_quotes:
                seen_quotes.add(quote_key)
                unique_insights.append(insight)
        
        return unique_insights
    
async def process(self, documents: List[Document]) -> List[InsightData]:
    """Process documents and extract insights"""
    try:
        all_insights = []
        
        for doc in documents:
            chunk_insights = await self._analyze_chunk(doc.page_content)
            all_insights.extend(chunk_insights)
        
        # Remove duplicates
        unique_insights = self._deduplicate_insights(all_insights)
        
        self.logger.info(f"Extracted {len(unique_insights)} unique insights from {len(documents)} chunks")
        self.update_session_status(
            ProcessingStatus.PROCESSING, 
            insights_extracted=len(unique_insights)
        )
        
        return unique_insights
        
    except Exception as e:
        self.logger.error(f"Error analyzing insights: {str(e)}")
        raise

async def _analyze_chunk(self, chunk: str) -> List[InsightData]:
    """Analyze a single chunk and extract insights"""
    try:
        # Create the prompt for this chunk
        formatted_prompt = self.insight_prompt.format(chunk=chunk)
        
        # Make the API call
        response = await self.llm.ainvoke(formatted_prompt)
        
        # Parse the response - expecting a list of insights
        insights_text = response.content
        
        # Basic parsing - you may need to adjust this based on your prompt format
        insights = []
        lines = insights_text.strip().split('\n')
        
        current_insight = {}
        for line in lines:
            line = line.strip()
            if line.startswith('Quote:'):
                current_insight['quote'] = line.replace('Quote:', '').strip()
            elif line.startswith('Theme:'):
                current_insight['theme'] = line.replace('Theme:', '').strip()
            elif line.startswith('Sentiment:'):
                current_insight['sentiment'] = line.replace('Sentiment:', '').strip()
            elif line.startswith('Confidence:'):
                try:
                    current_insight['confidence'] = float(line.replace('Confidence:', '').strip())
                except:
                    current_insight['confidence'] = 0.5
            elif line.startswith('Context:'):
                current_insight['context'] = line.replace('Context:', '').strip()
                # Complete insight, add to list
                if 'quote' in current_insight and 'theme' in current_insight:
                    insight_data = InsightData(
                        quote=current_insight.get('quote', ''),
                        theme=current_insight.get('theme', 'General'),
                        sentiment=current_insight.get('sentiment', 'Neutral'),
                        confidence=current_insight.get('confidence', 0.5),
                        context=current_insight.get('context', ''),
                        speaker=current_insight.get('speaker'),
                        timestamp=current_insight.get('timestamp')
                    )
                    insights.append(insight_data)
                current_insight = {}
        
        return insights
        
    except Exception as e:
        self.logger.error(f"Error analyzing chunk: {str(e)}")
        return []

def _deduplicate_insights(self, insights: List[InsightData]) -> List[InsightData]:
    """Remove duplicate insights based on quote similarity"""
    unique_insights = []
    seen_quotes = set()
    
    for insight in insights:
        # Simple deduplication based on quote content
        quote_key = insight.quote.lower().strip()[:50]  # First 50 chars, normalized
        if quote_key not in seen_quotes:
            seen_quotes.add(quote_key)
            unique_insights.append(insight)
    
    return unique_insights
    
async def _analyze_chunk(self, document: Document) -> List[InsightData]:
    """Analyze a single document chunk"""
    try:
        # Create a simple prompt without complex parsing
        simple_prompt = f"""
Extract insights from this user research text. Return valid JSON only:

{document.page_content}

Format as JSON array:
[{{"quote": "exact quote", "speaker": "User", "theme": "Theme Name", "sentiment": "Positive", "confidence": 0.8, "context": "brief description"}}]
"""
        
        response = await self.llm.ainvoke([{"role": "user", "content": simple_prompt}])
        
        # Parse JSON response
        import json
        content = response.content.strip()
        if content.startswith('```'):
            content = content.split('```')[1].replace('json', '').strip()
        
        data = json.loads(content)
        
        insights = []
        for item in data:
            insight = InsightData(
                quote=item.get('quote', ''),
                speaker=item.get('speaker'),
                theme=item.get('theme', 'General'),
                sentiment=item.get('sentiment', 'Neutral'),
                confidence=float(item.get('confidence', 0.5)),
                context=item.get('context', ''),
                timestamp=None
            )
            insights.append(insight)
        
        return insights
        
    except Exception as e:
        self.logger.error(f"Error in _analyze_chunk: {e}")
        return []

# Agent 3: Theme Synthesizer
class ThemeSynthesizer(BaseAgent):
    """Groups insights and generates personas/summaries"""
    
    def __init__(self, config: AgentConfig, session_id: str):
        super().__init__(config, session_id)
        
        if not config.openai_api_key:
            raise ValueError("OpenAI API key is required for ThemeSynthesizer")
            
        self.llm = ChatOpenAI(
            api_key=config.openai_api_key,
            model="gpt-3.5-turbo",
            temperature=config.temperature
        )
        
        try:
            # Use OpenAI embeddings (lightweight, no PyTorch needed)
            self.embedding_model = OpenAIEmbeddings(
                api_key=config.openai_api_key,
                model="text-embedding-3-small"  # Smaller, faster, cheaper
            )
        except Exception as e:
            self.logger.warning(f"Failed to load embedding model: {e}")
            self.embedding_model = None
            
        self._setup_prompts()
    
    def _setup_prompts(self):
        """Setup prompts for theme synthesis"""
        self.theme_prompt = PromptTemplate(
            input_variables=["insights"],
            template="""
            As a senior UX researcher, analyze these related insights and create a comprehensive theme summary.
            
            Insights:
            {insights}
            
            Create a theme that includes:
            1. A clear, descriptive theme name
            2. Priority level (High/Medium/Low) based on frequency and impact
            3. A comprehensive summary of the theme
            
            Return your analysis as a JSON object with the structure:
            {{
                "theme_name": "Clear Theme Name",
                "priority": "High/Medium/Low",
                "summary": "Detailed summary of the theme and its implications"
            }}
            """
        )
        
        self.persona_prompt = PromptTemplate(
            input_variables=["insights", "themes"],
            template="""
            Based on the following insights and themes, create 2-3 distinct user personas.
            
            Insights (first 20):
            {insights}
            
            Themes:
            {themes}
            
            For each persona, provide:
            1. A realistic name
            2. Demographics (age, occupation, tech comfort level)
            3. Primary goals related to the product/service
            4. Key pain points and frustrations
            5. Behavioral patterns
            6. Representative quotes from the research
            
            Return as a JSON list of personas.
            """
        )
    
    async def process(self, insights: List[InsightData]) -> List[ThemeCluster]:
        """Process insights to generate themes only"""
        try:
            # Cluster insights by theme similarity
            theme_clusters = await self._cluster_insights(insights)
            
            self.update_session_status(
                ProcessingStatus.PROCESSING,
                themes_identified=len(theme_clusters)
            )
            
            return theme_clusters
            
        except Exception as e:
            self.logger.error(f"Error synthesizing themes: {str(e)}")
            raise
    
    async def _cluster_insights(self, insights: List[InsightData]) -> List[ThemeCluster]:
        """Cluster insights by theme similarity"""
        if not insights:
            return []
        
        # Group insights by theme name
        theme_groups = {}
        for insight in insights:
            theme_name = insight.theme
            if theme_name not in theme_groups:
                theme_groups[theme_name] = []
            theme_groups[theme_name].append(insight)
        
        # Create theme clusters
        theme_clusters = []
        for theme_name, theme_insights in theme_groups.items():
            if len(theme_insights) >= 1:
                cluster = await self._create_theme_cluster(theme_name, theme_insights)
                theme_clusters.append(cluster)
        
        # Sort by frequency and priority
        theme_clusters.sort(key=lambda x: (x.priority == "High", x.frequency), reverse=True)
        
        return theme_clusters
    
    async def _create_theme_cluster(self, theme_name: str, insights: List[InsightData]) -> ThemeCluster:
        """Create a comprehensive theme cluster"""
        insights_text = "\n".join([
            f"- \"{insight.quote}\" (Sentiment: {insight.sentiment}, Confidence: {insight.confidence})"
            for insight in insights
        ])
        
        try:
            chain = self.theme_prompt | self.llm
            result = await chain.ainvoke({"insights": insights_text})
            
            # Parse the JSON response
            theme_data = json.loads(result.content)
            
            # Determine priority based on frequency and sentiment
            negative_count = sum(1 for insight in insights if insight.sentiment == "Negative")
            frequency = len(insights)
            
            if frequency >= 5 or (negative_count >= 2 and frequency >= 3):
                priority = "High"
            elif frequency >= 3:
                priority = "Medium"
            else:
                priority = "Low"
            
            return ThemeCluster(
                theme_name=theme_data.get("theme_name", theme_name),
                insights=insights,
                frequency=frequency,
                priority=priority,
                summary=theme_data.get("summary", f"Theme based on {frequency} insights")
            )
            
        except Exception as e:
            self.logger.warning(f"Error creating theme cluster: {str(e)}")
            # Fallback to basic cluster
            return ThemeCluster(
                theme_name=theme_name,
                insights=insights,
                frequency=len(insights),
                priority="Medium",
                summary=f"Theme identified from {len(insights)} user insights"
            )


# Agent 3.5: Key Insight Synthesizer
class KeyInsightSynthesizer(BaseAgent):
    """Generates Key Insight cards matching presentation format"""

    def __init__(self, config: AgentConfig, session_id: str):
        super().__init__(config, session_id)

        if not config.openai_api_key:
            raise ValueError("OpenAI API key is required for KeyInsightSynthesizer")

        self.llm = ChatOpenAI(
            api_key=config.openai_api_key,
            model="gpt-3.5-turbo-16k",  # Use larger context for better synthesis
            temperature=0.3  # Lower temperature for more focused output
        )
        self._setup_prompts()

    def _setup_prompts(self):
        """Setup prompts for key insight generation"""
        self.insight_synthesis_prompt = PromptTemplate(
            input_variables=["all_quotes", "theme_summaries"],
            template="""
            As an expert UX researcher, analyze all the user research data and create 3-5 Key Insight cards for a presentation.

            USER QUOTES AND FEEDBACK:
            {all_quotes}

            THEME SUMMARIES:
            {theme_summaries}

            Create 3-5 Key Insight cards. Each card should:
            1. Have a clear, compelling title (e.g., "Control & Transparency", "Privacy & Information Concerns")
            2. Include ONE main finding statement (what users want/feel/do)
            3. Specify the finding type: "positive" (users like), "negative" (users dislike), "critical" (must-have), or "neutral"
            4. Optionally include a problem statement explaining the issue
            5. Include 2-4 supporting quotes with speaker attribution
            6. Optionally include behavioral patterns observed
            7. Optionally include expected user journey steps
            8. Include impact metrics when quantifiable (e.g., "9 out of 11 participants")

            Return as a JSON array with this structure:
            [
              {{
                "insight_number": 1,
                "title": "Control & Transparency",
                "main_finding": "Users want visible choice upfront",
                "finding_type": "positive",
                "problem_statement": "Users had to click through to discover they had options, creating a hidden second decision point.",
                "supporting_quotes": [
                  {{"quote": "I like the one that on the home page makes it clear I had that option right here.", "speaker": "Bruce C., 69"}},
                  {{"quote": "It saves me a step. If I can only click on one button, then go to a two selection screen, I have to make that choice again.", "speaker": "Michael F., 64"}}
                ],
                "behavioral_pattern": "9 out of 11 participants said dual-button would make them more likely to complete the process",
                "impact_metric": "100% of participants preferred this approach"
              }}
            ]

            Focus on insights that are:
            - Actionable and specific
            - Supported by strong quotes
            - Relevant to product decisions
            - Clearly differentiated from each other
            """
        )

        self.executive_summary_prompt = PromptTemplate(
            input_variables=["key_insights", "all_data"],
            template="""
            As a senior UX researcher, create an Executive Summary for a research presentation.

            KEY INSIGHTS:
            {key_insights}

            ALL RESEARCH DATA:
            {all_data}

            Create an executive summary with:
            1. Research Question: The primary question this research aimed to answer
            2. Key Finding: The single most important discovery (be specific, include numbers if available)
            3. Key Insight: The interpretation/meaning of that finding
            4. Recommendation: The primary action stakeholders should take

            Return as JSON:
            {{
              "research_question": "The main research question...",
              "key_finding": "X out of Y participants preferred/did/said...",
              "key_insight": "This means that users...",
              "recommendation": "Implement/Change/Prioritize X to achieve Y.",
              "context": "Optional additional context paragraph"
            }}

            Be specific, quantitative when possible, and action-oriented.
            """
        )

    async def process(self, insights: List[InsightData], themes: List[ThemeCluster]) -> Tuple[List[KeyInsightCard], ExecutiveSummary]:
        """Generate Key Insight cards and Executive Summary"""
        try:
            # Prepare data for synthesis
            all_quotes = self._prepare_quotes(insights)
            theme_summaries = self._prepare_theme_summaries(themes)

            # Generate key insight cards
            key_insights = await self._generate_key_insights(all_quotes, theme_summaries)

            # Generate executive summary
            exec_summary = await self._generate_executive_summary(key_insights, all_quotes)

            self.logger.info(f"Generated {len(key_insights)} key insight cards")

            return key_insights, exec_summary

        except Exception as e:
            self.logger.error(f"Error synthesizing key insights: {str(e)}")
            raise

    def _prepare_quotes(self, insights: List[InsightData]) -> str:
        """Prepare all quotes for synthesis"""
        quotes_text = []
        for i, insight in enumerate(insights[:50], 1):  # Limit to top 50 for context
            speaker_info = f" - {insight.speaker}" if insight.speaker else ""
            quotes_text.append(
                f"{i}. \"{insight.quote}\"{speaker_info} "
                f"[Theme: {insight.theme}, Sentiment: {insight.sentiment}]"
            )
        return "\n".join(quotes_text)

    def _prepare_theme_summaries(self, themes: List[ThemeCluster]) -> str:
        """Prepare theme summaries"""
        summaries = []
        for theme in themes:
            summaries.append(
                f"- {theme.theme_name} ({theme.priority} priority, {theme.frequency} mentions): "
                f"{theme.summary}"
            )
        return "\n".join(summaries)

    async def _generate_key_insights(self, all_quotes: str, theme_summaries: str) -> List[KeyInsightCard]:
        """Generate key insight cards using LLM"""
        try:
            chain = self.insight_synthesis_prompt | self.llm
            result = await chain.ainvoke({
                "all_quotes": all_quotes,
                "theme_summaries": theme_summaries
            })

            # Parse JSON response
            insights_data = json.loads(result.content)

            # Convert to KeyInsightCard objects
            key_insights = []
            for i, insight_dict in enumerate(insights_data, 1):
                # Convert supporting quotes
                quotes = [
                    QuoteWithAttribution(**q) if isinstance(q, dict) else QuoteWithAttribution(quote=str(q))
                    for q in insight_dict.get("supporting_quotes", [])
                ]

                key_insight = KeyInsightCard(
                    insight_number=i,
                    title=insight_dict.get("title", f"Insight #{i}"),
                    main_finding=insight_dict.get("main_finding", ""),
                    finding_type=insight_dict.get("finding_type", "neutral"),
                    problem_statement=insight_dict.get("problem_statement"),
                    supporting_quotes=quotes,
                    behavioral_pattern=insight_dict.get("behavioral_pattern"),
                    expected_journey=insight_dict.get("expected_journey"),
                    impact_metric=insight_dict.get("impact_metric")
                )
                key_insights.append(key_insight)

            return key_insights

        except Exception as e:
            self.logger.error(f"Error generating key insights: {str(e)}")
            # Return empty list on error
            return []

    async def _generate_executive_summary(self, key_insights: List[KeyInsightCard], all_data: str) -> ExecutiveSummary:
        """Generate executive summary"""
        try:
            # Prepare key insights summary
            insights_summary = "\n".join([
                f"{i.insight_number}. {i.title}: {i.main_finding}"
                for i in key_insights
            ])

            chain = self.executive_summary_prompt | self.llm
            result = await chain.ainvoke({
                "key_insights": insights_summary,
                "all_data": all_data[:3000]  # Limit for context
            })

            # Parse JSON response
            summary_data = json.loads(result.content)

            return ExecutiveSummary(**summary_data)

        except Exception as e:
            self.logger.error(f"Error generating executive summary: {str(e)}")
            # Return a default summary
            return ExecutiveSummary(
                research_question="User research analysis",
                key_finding=f"{len(key_insights)} key insights identified from user research",
                key_insight="Multiple themes emerged from the research",
                recommendation="Review key insights and prioritize implementation",
                context=None
            )


# Agent 4: Output Formatter
class OutputFormatter(BaseAgent):
    """Creates structured deliverables for stakeholders"""
    
    def __init__(self, config: AgentConfig, session_id: str):
        super().__init__(config, session_id)
        self.output_dir = Path(f"outputs/{session_id}")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    async def process(self, insights: List[InsightData], themes: List[ThemeCluster],
                     personas: List[PersonaData]) -> Dict[str, str]:
        """Generate formatted outputs"""
        try:
            outputs = {}

            # Generate JSON report
            json_path = await self._create_json_report(insights, themes, personas)
            outputs["json_report"] = str(json_path)

            # Generate executive summary
            summary_path = await self._create_executive_summary(themes, personas)
            outputs["executive_summary"] = str(summary_path)

            # Generate detailed insights report
            insights_path = await self._create_insights_report(insights, themes)
            outputs["insights_report"] = str(insights_path)

            # Generate persona profiles
            personas_path = await self._create_persona_profiles(personas)
            outputs["persona_profiles"] = str(personas_path)

            self.update_session_status(ProcessingStatus.COMPLETED)

            return outputs

        except Exception as e:
            self.logger.error(f"Error formatting outputs: {str(e)}")
            self.update_session_status(ProcessingStatus.FAILED, error_message=str(e))
            raise

    async def process_with_key_insights(self, insights: List[InsightData], themes: List[ThemeCluster],
                                       key_insights: List[KeyInsightCard], exec_summary: ExecutiveSummary) -> Dict[str, str]:
        """Generate formatted outputs with key insights"""
        try:
            outputs = {}

            # Generate JSON report with key insights
            json_path = await self._create_json_report_with_key_insights(insights, themes, key_insights, exec_summary)
            outputs["json_report"] = str(json_path)

            # Generate executive summary (new format)
            summary_path = await self._create_executive_summary_new(exec_summary, key_insights)
            outputs["executive_summary"] = str(summary_path)

            # Generate detailed insights report
            insights_path = await self._create_insights_report(insights, themes)
            outputs["insights_report"] = str(insights_path)

            # Generate persona profiles (empty for now)
            personas_path = await self._create_persona_profiles([])
            outputs["persona_profiles"] = str(personas_path)

            self.update_session_status(ProcessingStatus.COMPLETED)

            return outputs

        except Exception as e:
            self.logger.error(f"Error formatting outputs: {str(e)}")
            self.update_session_status(ProcessingStatus.FAILED, error_message=str(e))
            raise

    async def _create_json_report_with_key_insights(self, insights: List[InsightData], themes: List[ThemeCluster],
                                                   key_insights: List[KeyInsightCard], exec_summary: ExecutiveSummary) -> Path:
        """Create comprehensive JSON report with key insights"""
        report = {
            "session_id": self.session_id,
            "generated_at": datetime.now().isoformat(),
            "executive_summary": {
                "research_question": exec_summary.research_question,
                "key_finding": exec_summary.key_finding,
                "key_insight": exec_summary.key_insight,
                "recommendation": exec_summary.recommendation,
                "context": exec_summary.context
            },
            "key_insights": [
                {
                    "insight_number": ki.insight_number,
                    "title": ki.title,
                    "main_finding": ki.main_finding,
                    "finding_type": ki.finding_type,
                    "problem_statement": ki.problem_statement,
                    "supporting_quotes": [
                        {"quote": q.quote, "speaker": q.speaker, "context": q.context}
                        for q in ki.supporting_quotes
                    ],
                    "behavioral_pattern": ki.behavioral_pattern,
                    "expected_journey": ki.expected_journey,
                    "impact_metric": ki.impact_metric
                }
                for ki in key_insights
            ],
            "summary": {
                "total_insights": len(insights),
                "themes_identified": len(themes),
                "key_insights_count": len(key_insights),
                "personas_created": 0,
                "quotes_extracted": sum(len(ki.supporting_quotes) for ki in key_insights),
                "files_processed": 5,  # This should come from session data
                "processing_time": "2m 34s",  # This should be calculated
                "total_tokens": sum([len(i.quote.split()) for i in insights]) * 1.3  # Rough estimate
            },
            "themes": [
                {
                    "theme_name": theme.theme_name,
                    "frequency": theme.frequency,
                    "priority": theme.priority,
                    "summary": theme.summary,
                    "insights": [insight.dict() for insight in theme.insights]
                }
                for theme in themes
            ],
            "insights": [insight.dict() for insight in insights]
        }

        json_path = self.output_dir / "research_synthesis.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)

        return json_path

    async def _create_executive_summary_new(self, exec_summary: ExecutiveSummary, key_insights: List[KeyInsightCard]) -> Path:
        """Create executive summary with new presentation format"""
        summary_content = f"""# Executive Summary

**Session ID:** {self.session_id}
**Generated:** {datetime.now().strftime("%Y-%m-%d %H:%M")}

## Research Question

{exec_summary.research_question}

## Key Finding

**{exec_summary.key_finding}**

## Key Insight

{exec_summary.key_insight}

## Recommendation

{exec_summary.recommendation}

---

## Key Insights Overview

"""
        for ki in key_insights:
            summary_content += f"""
### {ki.insight_number}. {ki.title}

**{ki.main_finding}**

"""
            if ki.impact_metric:
                summary_content += f"*{ki.impact_metric}*\n\n"

        summary_path = self.output_dir / "executive_summary.md"
        with open(summary_path, 'w', encoding='utf-8') as f:
            f.write(summary_content)

        return summary_path
    
    async def _create_json_report(self, insights: List[InsightData], themes: List[ThemeCluster], 
                                 personas: List[PersonaData]) -> Path:
        """Create comprehensive JSON report"""
        report = {
            "session_id": self.session_id,
            "generated_at": datetime.now().isoformat(),
            "summary": {
                "total_insights": len(insights),
                "themes_identified": len(themes),
                "personas_created": len(personas)
            },
            "insights": [insight.dict() for insight in insights],
            "themes": [
                {
                    "theme_name": theme.theme_name,
                    "frequency": theme.frequency,
                    "priority": theme.priority,
                    "summary": theme.summary,
                    "insights": [insight.dict() for insight in theme.insights]
                }
                for theme in themes
            ],
            "personas": [persona.dict() for persona in personas]
        }
        
        json_path = self.output_dir / "research_synthesis.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        
        return json_path
    
    async def _create_executive_summary(self, themes: List[ThemeCluster], 
                                      personas: List[PersonaData]) -> Path:
        """Create executive summary document"""
        summary_content = f"""# UX Research Synthesis - Executive Summary

**Session ID:** {self.session_id}
**Generated:** {datetime.now().strftime("%Y-%m-%d %H:%M")}

## Key Findings

### Top Themes ({len(themes)} identified)
"""
        
        for theme in themes[:5]:  # Top 5 themes
            summary_content += f"""
#### {theme.theme_name} ({theme.priority} Priority)
- **Frequency:** {theme.frequency} mentions
- **Summary:** {theme.summary}
"""
        
        summary_content += f"""
### User Personas ({len(personas)} created)
"""
        
        for persona in personas:
            summary_content += f"""
#### {persona.name}
- **Demographics:** {persona.demographics}
- **Top Goals:** {', '.join(persona.goals[:3])}
- **Main Pain Points:** {', '.join(persona.pain_points[:3])}
"""
        
        summary_content += """
## Recommendations

1. **Immediate Actions:** Address high-priority themes with multiple negative sentiments
2. **Design Focus:** Consider persona needs in upcoming design decisions
3. **Further Research:** Investigate themes with medium priority for deeper understanding

---

*This summary was automatically generated by the UX Research Copilot system.*
"""
        
        summary_path = self.output_dir / "executive_summary.md"
        with open(summary_path, 'w', encoding='utf-8') as f:
            f.write(summary_content)
        
        return summary_path
    
    async def _create_insights_report(self, insights: List[InsightData], 
                                    themes: List[ThemeCluster]) -> Path:
        """Create detailed insights report"""
        report_content = f"""# Detailed Insights Report

**Session ID:** {self.session_id}
**Total Insights:** {len(insights)}
**Generated:** {datetime.now().strftime("%Y-%m-%d %H:%M")}

## Insights by Theme
"""
        
        for theme in themes:
            report_content += f"""
### {theme.theme_name} ({theme.frequency} insights)
**Priority:** {theme.priority}
**Summary:** {theme.summary}

**Key Quotes:**
"""
            for insight in theme.insights:
                report_content += f"""
- "{insight.quote}"
  - **Sentiment:** {insight.sentiment}
  - **Confidence:** {insight.confidence:.2f}
  - **Context:** {insight.context}
"""
        
        insights_path = self.output_dir / "detailed_insights.md"
        with open(insights_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        return insights_path
    
    async def _create_persona_profiles(self, personas: List[PersonaData]) -> Path:
        """Create detailed persona profiles"""
        profiles_content = f"""# User Persona Profiles

**Session ID:** {self.session_id}
**Generated:** {datetime.now().strftime("%Y-%m-%d %H:%M")}

"""
        
        for i, persona in enumerate(personas, 1):
            profiles_content += f"""
## Persona {i}: {persona.name}

**Demographics:** {persona.demographics}

### Goals
"""
            for goal in persona.goals:
                profiles_content += f"- {goal}\n"
            
            profiles_content += """
### Pain Points
"""
            for pain_point in persona.pain_points:
                profiles_content += f"- {pain_point}\n"
            
            profiles_content += """
### Behaviors
"""
            for behavior in persona.behaviors:
                profiles_content += f"- {behavior}\n"
            
            profiles_content += """
### Representative Quotes
"""
            for quote in persona.quotes:
                profiles_content += f'- "{quote}"\n'
            
            profiles_content += "\n---\n"
        
        personas_path = self.output_dir / "persona_profiles.md"
        with open(personas_path, 'w', encoding='utf-8') as f:
            f.write(profiles_content)
        
        return personas_path

# Multi-Agent Orchestrator
class UXResearchOrchestrator:
    """Main orchestrator that coordinates all agents"""
    
    def __init__(self, config: AgentConfig):
        self.config = config
        self.logger = logging.getLogger("UXResearchOrchestrator")
    
    async def process_research_files_with_session(self, file_paths: List[str], session_id: str) -> Dict[str, Any]:
        """Complete research processing workflow"""
        
        # Initialize Redis client
        redis_client = None
        try:
            redis_client = redis.Redis(
                host=self.config.redis_host,
                port=self.config.redis_port,
                decode_responses=True
            )
            redis_client.ping()
            self.logger.info("Redis connection established")
        except:
            self.logger.warning("Redis not available - continuing without session persistence")
        
        try:
            # Initialize session
            if redis_client:
                redis_client.hmset(f"session:{session_id}", {
                    'session_id': session_id,
                    'status': 'processing',
                    'current_phase': 'initialization',
                    'created_at': datetime.now().isoformat(),
                    'updated_at': datetime.now().isoformat(),
                    'file_count': len(file_paths),
                    'insights_extracted': 0,
                    'themes_identified': 0,
                    'personas_created': 0
                })
            
            self.logger.info(f"Starting processing for session {session_id}")
            
            # Phase 1: Document Ingestion
            self.logger.info("Phase 1: Document Ingestion")
            if redis_client:
                redis_client.hset(f"session:{session_id}", 'current_phase', 'document_ingestion')
            
            ingestor = DocumentIngestor(self.config, session_id)
            documents = await ingestor.process(file_paths)
            
            if not documents:
                raise ValueError("No documents were successfully processed")
            
            self.logger.info(f"Phase 1 complete: {len(documents)} documents processed")
            
            # Phase 2: Insight Analysis
            self.logger.info("Phase 2: Insight Analysis")
            if redis_client:
                redis_client.hset(f"session:{session_id}", 'current_phase', 'insight_analysis')
            
            analyzer = InsightAnalyzer(self.config, session_id)
            insights = await analyzer.process(documents)
            
            # Update Redis with insights count
            if redis_client:
                redis_client.hmset(f"session:{session_id}", {
                    'insights_extracted': len(insights),
                    'updated_at': datetime.now().isoformat()
                })
            
            self.logger.info(f"Phase 2 complete: {len(insights)} insights extracted")
            
            if not insights:
                self.logger.warning("No insights extracted - creating minimal results")
                insights = [InsightData(
                    quote="No significant insights found in the provided documents",
                    speaker=None,
                    theme="Analysis Result",
                    sentiment="Neutral",
                    confidence=1.0,
                    context="System message",
                    timestamp=None
                )]
            
            # Phase 3: Theme Synthesis
            self.logger.info("Phase 3: Theme Synthesis")
            if redis_client:
                redis_client.hset(f"session:{session_id}", 'current_phase', 'theme_synthesis')

            synthesizer = ThemeSynthesizer(self.config, session_id)
            themes = await synthesizer.process(insights)

            # Update Redis with themes count only
            if redis_client:
                redis_client.hmset(f"session:{session_id}", {
                'themes_identified': len(themes),
                'updated_at': datetime.now().isoformat()
                })

            self.logger.info(f"Phase 3 complete: {len(themes)} themes")

            # Phase 3.5: Key Insight Synthesis (NEW)
            self.logger.info("Phase 3.5: Key Insight Synthesis")
            if redis_client:
                redis_client.hset(f"session:{session_id}", 'current_phase', 'key_insight_synthesis')

            key_insight_synthesizer = KeyInsightSynthesizer(self.config, session_id)
            key_insights, executive_summary = await key_insight_synthesizer.process(insights, themes)

            self.logger.info(f"Phase 3.5 complete: {len(key_insights)} key insight cards generated")

            # Phase 4: Output Formatting
            self.logger.info("Phase 4: Output Formatting")
            if redis_client:
                redis_client.hset(f"session:{session_id}", 'current_phase', 'output_formatting')

            formatter = OutputFormatter(self.config, session_id)
            # Pass key insights and executive summary to formatter
            outputs = await formatter.process_with_key_insights(insights, themes, key_insights, executive_summary)
            
            self.logger.info(f"Phase 4 complete: {len(outputs)} output files created")
            
            # Verify files were created
            for output_type, file_path in outputs.items():
                if os.path.exists(file_path):
                    file_size = os.path.getsize(file_path)
                    self.logger.info(f"✓ {output_type}: {file_path} ({file_size} bytes)")
                else:
                    self.logger.error(f"✗ {output_type}: {file_path} - FILE NOT FOUND")
            
            # Mark as completed
            if redis_client:
                redis_client.hmset(f"session:{session_id}", {
                    'status': 'completed',
                    'current_phase': 'completed',
                    'updated_at': datetime.now().isoformat()
                })
            
            # Create comprehensive results
            results = {
                "session_id": session_id,
                "status": "completed",
                "results": {
                    "insights_count": len(insights),
                    "themes_count": len(themes),
                    "personas_count": 0
                },
                "outputs": outputs,
                "insights": [insight.dict() for insight in insights],
                "themes": [
                    {
                        "theme_name": theme.theme_name,
                        "frequency": theme.frequency,
                        "priority": theme.priority,
                        "summary": theme.summary,
                        "insights": [insight.dict() for insight in theme.insights]
                    }
                    for theme in themes
                ],
                "personas": []
            }
            
            self.logger.info(f"Processing completed successfully for session {session_id}")
            return results
            
        except Exception as e:
            self.logger.error(f"Error in orchestration: {str(e)}")
            
            # Mark as failed in Redis
            if redis_client:
                redis_client.hmset(f"session:{session_id}", {
                    'status': 'failed',
                    'error_message': str(e),
                    'updated_at': datetime.now().isoformat()
                })
            
            raise

# Example usage and testing functions
async def test_orchestrator():
    """Test function for the orchestrator"""
    from config import CONFIG
    
    # Example file paths - replace with actual paths
    sample_files = [
        "sample_data/interview_1.txt",
        "sample_data/interview_2.txt"
    ]
    
    try:
        orchestrator = UXResearchOrchestrator(CONFIG)
        results = await orchestrator.process_research_files(sample_files)
        
        print(f"✅ Processing completed!")
        print(f"📊 Session ID: {results['session_id']}")
        print(f"⏱️ Processing time: {results['processing_time']:.2f} seconds")
        print(f"🔍 Insights extracted: {results['results']['insights_count']}")
        print(f"🎯 Themes identified: {results['results']['themes_count']}")
        print(f"👥 Personas created: {results['results']['personas_count']}")
        
        return results
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return None


if __name__ == "__main__":
    """
    Run this file directly to test the agents
    Usage: python agents.py
    """
    import asyncio
    
    print("🧪 Testing UX Research Copilot Agents...")
    
    # Check if configuration is valid
    from config import CONFIG
    
    if not CONFIG.openai_api_key or CONFIG.openai_api_key == "your_openai_api_key_here":
        print("❌ Please configure your OpenAI API key in the .env file")
        exit(1)
    
    # Run the test
    asyncio.run(test_orchestrator())